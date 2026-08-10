from copy import deepcopy
from pathlib import Path
from docx import Document


doc_path = next(Path('.').glob('*PU3更新版.docx'))
document = Document(doc_path)

# 删除先前插入的普通段落，避免与 X 轴章节的标题/表格结构不一致。
for paragraph in list(document.paragraphs):
    if paragraph.text.startswith('1E：PU3') or paragraph.text.startswith('AA 1E') or paragraph.text.startswith('STEP_H STEP_L 为大端'):
        paragraph._element.getparent().remove(paragraph._element)

x_heading = next(p for p in document.paragraphs if p.text == '1B：X 轴 PB11')
y_heading = next(p for p in document.paragraphs if p.text == '1C：Y 轴 PB10')

# 复用 X 轴 Heading 2 样式，并插入到 Y 轴之后。
z_heading = y_heading.insert_paragraph_before('1E：PU3 轴 PB13', style=x_heading.style)
y_heading._p.addnext(z_heading._p)

# 复制 X 轴的一行帧表，替换命令号，确保字体、边框和宽度一致。
x_table = next(t for t in document.tables if t.cell(0, 0).text.startswith('AA 1B'))
z_table_xml = deepcopy(x_table._tbl)
for node in z_table_xml.iter():
    if node.tag.endswith('}t') and node.text:
        node.text = node.text.replace('1B', '1E')
z_heading._p.addnext(z_table_xml)

explanation = document.add_paragraph('STEP_H STEP_L 为大端 16 位步数，DIR 仅使用 bit0。PU3 起步速度为 500 steps/s，最大速度为 SPEED * 200 steps/s，加速度为 100000 steps/s²。STEP 由 TIM7 更新事件触发 DMA2 Channel4 向 GPIOB->BSRR 交替写 PB13 的置位和复位值；两个 DMA edge 构成一个实际步数。DIR 为 U86（第二片 74HC595）Q6、软件 bit6。')
z_table_xml.addnext(explanation._p)

document.save(doc_path.with_name(doc_path.stem + '_排版修正.docx'))
