from pathlib import Path
from docx import Document


def add_after(paragraph, text):
    item = paragraph.insert_paragraph_before(text)
    paragraph._p.addprevious(item._p)
    return item


source = next(Path('.').glob('*.docx'))
document = Document(source)

for paragraph in document.paragraphs:
    if paragraph.text == '1C：Y 轴 PB10':
        add_after(paragraph, '1E：PU3 轴 PB13')
        add_after(paragraph, 'AA 1E STEP_H STEP_L DIR SPEED 00 00 55 SUM')
        add_after(paragraph, 'STEP_H STEP_L 为大端 16 位步数，DIR 仅使用 bit0。PU3 起步速度为 500 steps/s，最大速度为 SPEED * 200 steps/s，加速度为 100000 steps/s²。STEP 由 TIM7 更新事件触发 DMA2 Channel4 向 GPIOB->BSRR 交替写 PB13 的置位和复位值；两个 DMA edge 构成一个实际步数。DIR 为 U86（第二片 74HC595）Q6、软件 bit6。')
        break

for table in document.tables:
    for row in table.rows:
        if row.cells[0].text == '1C':
            new_row = table.add_row()
            new_row.cells[0].text = '1E'
            new_row.cells[1].text = 'PU3 步进电机（PB13）'
            new_row.cells[2].text = '原样回显'
            break

document.save(source.with_name(source.stem + '_PU3更新版.docx'))
