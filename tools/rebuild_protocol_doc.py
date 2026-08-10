from pathlib import Path
from docx import Document


template = next(p for p in Path('.').glob('*.docx') if p.stat().st_size == 34002)
doc = Document(template)
body = doc._body._element
for child in list(body)[:-1]:
    body.remove(child)


def frame(cmd, fields):
    table = doc.add_table(rows=3, cols=10, style='Normal Table')
    headers = ['STX', 'CMD', 'DATA1', 'DATA2', 'DATA3', 'DATA4', 'DATA5', 'DATA6', 'ETX', 'SUM']
    values = ['AAH', cmd + 'H'] + fields + ['55H', 'SUM']
    descriptions = {
        '00': '预留', 'MOTOR': '电机编号', 'ACTION': '动作/方向', 'DIR': '方向',
        'TIME_L': '时间低8位', 'TIME_H': '时间次低8位', 'TIME_HH': '时间次高8位',
        'TIME_HHH': '时间高8位', 'TIME100_L': '超时低8位', 'TIME100_H': '超时高8位',
        'SENSOR': '传感器编号', 'LEVEL': '触发电平', 'RESULT': '结果码',
        'MAP0': '电机1..8位图', 'MAP1': '电机9..16位图', 'MAP2': '电机17..19位图',
        'STATE1': '电机1状态', 'STATE2': '电机2状态', 'STATE3': '电机3状态',
        'STATE4_5': '电机4/5状态', 'REASON1': '电机1停止原因', 'REASON2': '电机2停止原因',
        'STEP_H': '步数高8位', 'STEP_L': '步数低8位', 'SPEED': '最大速度系数',
        'STATUS': '运行状态', 'D0': '数据0', 'D1': '数据1'
    }
    meanings = ['开始码', '命令'] + [descriptions.get(value, '预留') for value in fields] + ['结束码', '校验和']
    for index, value in enumerate(headers): table.rows[0].cells[index].text = value
    for index, value in enumerate(values): table.rows[1].cells[index].text = value
    for index, value in enumerate(meanings): table.rows[2].cells[index].text = value


def command(index, cmd, name, request, reply, note=''):
    doc.add_paragraph(f'{index})  当前协议  CMD={cmd}: {name}', style='Heading 2')
    doc.add_paragraph('主机发送')
    frame(cmd, request)
    doc.add_paragraph('从机回应')
    frame(cmd, reply)
    if note:
        doc.add_paragraph(note)


doc.add_paragraph('蛋白粉机串口协议文档（当前代码版）', style='Title')
doc.add_paragraph('交互格式', style='Heading 1')
doc.add_paragraph('串口参数：USART2，9600bps，8位数据，1位停止位，无校验，无硬件流控。帧格式固定为 AA、CMD、6 个数据字节、55、SUM；SUM 为前 8 字节累加后的低 8 位。')
doc.add_paragraph('通讯格式', style='Heading 1')

commands = [
('00','保留空操作',['00']*6,['00']*6,'当前命令不产生业务动作。'),
('01','读取 74HC165 输入状态',['00']*6,['D0','D1','00','00','00','00'],'D0..D1 为当前 16 路输入快照。'),
('05','19 路单向电机立即控制',['MOTOR','ACTION','00','00','00','00'],['RESULT','00','00','00','00','00'],'MOTOR 为 1..19；ACTION 为 00 停止或 01 启动。'),
('06','19 路单向电机定时运行',['MOTOR','TIME_L','TIME_H','TIME_HH','TIME_HHH','00'],['RESULT','00','00','00','00','00'],'时间为小端 32 位毫秒，必须大于 0。'),
('07','19 路单向电机传感器或超时停止',['MOTOR','SENSOR','LEVEL','TIME100_L','TIME100_H','00'],['RESULT','00','00','00','00','00'],'SENSOR 为 1..16；超时单位 100ms。'),
('08','查询 19 路单向电机状态',['00']*6,['MAP0','MAP1','MAP2','00','00','00'],'MAP0..MAP2 为 19 路运行位图。'),
('09','五路正反转电机立即控制',['MOTOR','ACTION','00','00','00','00'],['RESULT','00','00','00','00','00'],'MOTOR 为 1..5；ACTION 为 00 停止、01 正转、02 反转。'),
('0A','五路正反转电机定时运行',['MOTOR','DIR','TIME_L','TIME_H','TIME_HH','TIME_HHH'],['RESULT','00','00','00','00','00'],'时间为小端 32 位毫秒。'),
('0B','五路正反转电机传感器或超时停止',['MOTOR','DIR','SENSOR','LEVEL','TIME100_L','TIME100_H'],['RESULT','00','00','00','00','00'],'SENSOR 为 1..16；超时单位 100ms。'),
('0C','查询五路正反转电机状态',['00']*6,['STATE1','STATE2','REASON1','REASON2','STATE3','STATE4_5'],'状态 00 停止、01 正转、02 反转、03 换向死区。'),
('1B','PU1 / X 轴 PB11 步进电机',['STEP_H','STEP_L','DIR','SPEED','00','00'],['STEP_H','STEP_L','DIR','SPEED','00','00'],'STEP 为大端 16 位；DIR 使用 bit0；最大速度为 SPEED * 200 steps/s。'),
('1C','PU2 / Y 轴 PB10 步进电机',['STEP_H','STEP_L','DIR','SPEED','00','00'],['STEP_H','STEP_L','DIR','SPEED','00','00'],'STEP 为大端 16 位；DIR 使用 bit0；最大速度为 SPEED * 200 steps/s。'),
('1E','PU3 / Z 轴 PB13 步进电机',['STEP_H','STEP_L','DIR','SPEED','00','00'],['STEP_H','STEP_L','DIR','SPEED','00','00'],'PU3 使用 TIM7 + DMA2 Channel4 写 GPIOB->BSRR；DIR 为 U86 Q6、第二片 74HC595 bit6。'),
('1F','查询 PU2 步进状态',['00']*6,['STATUS','00','00','00','00','00'],'STATUS=00 运行中，01 已停止或完成。'),
('20','PU2 传感器停止步进',['STEP_H','STEP_L','DIR','SPEED','SENSOR','LEVEL'],['STEP_H','STEP_L','DIR','SPEED','SENSOR','LEVEL'],'前四个字节与 1C 一致；传感器触发后主动上报 1F。'),
]
for index, (cmd, name, request, reply, note) in enumerate(commands):
    command(index, cmd, name, request, reply, note)

doc.add_paragraph('结果码', style='Heading 1')
doc.add_paragraph('05..0B 的应答中 D0 为结果码：00 成功、01 电机编号非法、02 动作或保留参数非法、03 传感器编号非法、04 时间非法、05 命令不支持。')
doc.save('ProteinControl_串口协议说明_按示例格式_完整说明.docx')
